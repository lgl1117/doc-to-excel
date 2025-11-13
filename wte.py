import tkinter as tk
from tkinter import filedialog, ttk, messagebox
import os
import threading
from docx import Document
import pandas as pd
from openpyxl.styles import Font, Border, Side, Alignment, PatternFill
from openpyxl.utils import get_column_letter
import math

class WordTableExtractor:
    def __init__(self, root):
        self.root = root
        self.root.title("Word表格提取工具")
        self.root.geometry("600x450")
        self.root.resizable(False, False)
        
        # 设置窗口图标
        try:
            self.root.iconbitmap('tubiao.png')
        except:
            # 如果图标设置失败，不影响程序运行
            pass
        
        # 设置紫色系主题
        self.style = ttk.Style()
        
        # 配置主题颜色 - 使用纯白色背景和宋体字
        self.style.configure("TLabel", 
                            font=('SimSun', 11, 'normal'),
                            foreground='#5D3FD3',
                            background='#FFFFFF',
                            padding=5)
        
        self.style.configure("TButton", 
                            font=('SimSun', 10, 'bold'),
                            foreground='#5D3FD3',
                            background='#FFFFFF',
                            padding=5)
        
        self.style.map("TButton",
                      background=[('pressed', '#F0F0F0'), ('active', '#E6E6FA')],
                      foreground=[('pressed', '#5D3FD3'), ('active', '#5D3FD3')])
        
        self.style.configure("TEntry",
                            font=('SimSun', 10, 'normal'),
                            background='#FFFFFF',
                            padding=3)
        
        self.style.configure("TProgressbar",
                            troughcolor='#FFFFFF',
                            background='#9370DB')
        
        # 创建主框架，使用纯白色背景
        # 为ttk.Frame创建一个白色背景的样式
        self.style.configure("White.TFrame", background="#FFFFFF")
        main_frame = ttk.Frame(root, padding="30", style="White.TFrame")
        main_frame.pack(fill=tk.BOTH, expand=True)
        # 设置窗口背景为白色
        root.configure(bg="#FFFFFF")
        
        # 创建标题标签
        title_label = ttk.Label(main_frame, 
                               text="Word表格提取工具", 
                               font=('SimSun', 16, 'bold'),
                               foreground='#5D3FD3')
        title_label.grid(row=0, column=0, columnspan=3, pady=15)
        
        # 添加分隔线
        separator = ttk.Separator(main_frame, orient='horizontal')
        separator.grid(row=1, column=0, columnspan=3, sticky='ew', pady=10)
        
        # Word文件选择 - 调整为更美观的布局
        file_frame = ttk.Frame(main_frame, style="White.TFrame")
        file_frame.grid(row=2, column=0, columnspan=3, sticky='ew', pady=5)
        
        ttk.Label(file_frame, text="Word文件:", anchor='w').pack(side=tk.LEFT, padx=5, fill=tk.Y)
        
        file_entry_frame = ttk.Frame(file_frame, style="White.TFrame")
        file_entry_frame.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        self.word_path_var = tk.StringVar()
        word_entry = ttk.Entry(file_entry_frame, textvariable=self.word_path_var)
        word_entry.pack(fill=tk.X, expand=True, ipady=3)
        
        ttk.Button(file_frame, text="浏览...", command=self.select_word_file, width=10).pack(side=tk.RIGHT, padx=5)
        
        # Excel保存路径选择
        excel_frame = ttk.Frame(main_frame, style="White.TFrame")
        excel_frame.grid(row=3, column=0, columnspan=3, sticky='ew', pady=5)
        
        ttk.Label(excel_frame, text="Excel保存路径:", anchor='w').pack(side=tk.LEFT, padx=5, fill=tk.Y)
        
        excel_entry_frame = ttk.Frame(excel_frame, style="White.TFrame")
        excel_entry_frame.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        self.excel_path_var = tk.StringVar()
        excel_entry = ttk.Entry(excel_entry_frame, textvariable=self.excel_path_var)
        excel_entry.pack(fill=tk.X, expand=True, ipady=3)
        
        ttk.Button(excel_frame, text="浏览...", command=self.select_excel_file, width=10).pack(side=tk.RIGHT, padx=5)
        
        # 状态标签 - 使用更醒目的紫色
        self.status_var = tk.StringVar(value="就绪")
        status_frame = ttk.Frame(main_frame, style="White.TFrame")
        status_frame.grid(row=4, column=0, columnspan=3, pady=10, sticky='ew')
        
        ttk.Label(status_frame, text="状态:", anchor='w').pack(side=tk.LEFT)
        status_label = ttk.Label(status_frame, textvariable=self.status_var, foreground="#8B008B", font=('SimSun', 10, 'bold'))
        status_label.pack(side=tk.LEFT, padx=5)
        
        # 进度条 - 调整长度和样式
        progress_frame = ttk.Frame(main_frame, style="White.TFrame")
        progress_frame.grid(row=5, column=0, columnspan=3, pady=10, sticky='ew')
        
        ttk.Label(progress_frame, text="进度:", anchor='w').pack(side=tk.LEFT, padx=5, pady=5)
        
        self.progress = ttk.Progressbar(progress_frame, orient=tk.HORIZONTAL, length=500, mode='determinate')
        self.progress.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        # 按钮框架 - 确保按钮正确显示
        button_frame = ttk.Frame(main_frame, style="White.TFrame")
        button_frame.grid(row=6, column=0, columnspan=3, pady=20, sticky='ew')
        
        # 添加按钮容器，确保居中显示
        buttons_container = ttk.Frame(button_frame, style="White.TFrame")
        buttons_container.pack(anchor='center', pady=10)
        
        # 启动按钮
        self.extract_button = ttk.Button(buttons_container, text="开始转换", command=self.start_extraction, width=15)
        self.extract_button.pack(side=tk.LEFT, padx=20, ipady=5)
        
        # 退出按钮
        exit_button = ttk.Button(buttons_container, text="退出", command=root.quit, width=15)
        exit_button.pack(side=tk.LEFT, padx=20, ipady=5)
        

    
    def select_word_file(self):
        file_path = filedialog.askopenfilename(
            title="选择Word文档",
            filetypes=[("Word文件", "*.docx"), ("所有文件", "*.*")]
        )
        if file_path:
            self.word_path_var.set(file_path)
            # 自动设置Excel保存路径（与Word同目录）
            word_dir = os.path.dirname(file_path)
            word_name = os.path.splitext(os.path.basename(file_path))[0]
            default_excel_path = os.path.join(word_dir, f"{word_name}_提取表格.xlsx")
            self.excel_path_var.set(default_excel_path)
    
    def select_excel_file(self):
        file_path = filedialog.asksaveasfilename(
            title="保存Excel文件",
            defaultextension=".xlsx",
            filetypes=[("Excel文件", "*.xlsx"), ("所有文件", "*.*")]
        )
        if file_path:
            self.excel_path_var.set(file_path)
    
    def start_extraction(self):
        word_path = self.word_path_var.get()
        excel_path = self.excel_path_var.get()
        
        # 验证输入
        if not word_path or not excel_path:
            messagebox.showerror("错误", "请选择Word文件和Excel保存路径")
            return
        
        if not os.path.exists(word_path):
            messagebox.showerror("错误", "Word文件不存在")
            return
        
        # 禁用按钮，防止重复点击
        self.extract_button.config(state=tk.DISABLED)
        self.status_var.set("正在提取表格...")
        self.progress["value"] = 0
        
        # 在新线程中执行提取操作，避免界面冻结
        threading.Thread(target=self.extract_tables, args=(word_path, excel_path), daemon=True).start()
    
    def extract_tables(self, word_path, excel_path):
        try:
            # 打开Word文档
            doc = Document(word_path)
            total_tables = len(doc.tables)
            
            if total_tables == 0:
                self.root.after(0, lambda: messagebox.showinfo("提示", "Word文档中未发现表格"))
                self.root.after(0, self.reset_ui)
                return
            
            # 创建Excel写入器
            excel_writer = pd.ExcelWriter(excel_path, engine="openpyxl")
            
            # 遍历所有表格并写入Excel不同工作表
            for i, table in enumerate(doc.tables):
                # 更新进度
                progress_value = ((i + 1) / total_tables) * 100
                self.root.after(0, lambda val=progress_value: self.update_progress(val))
                self.root.after(0, lambda num=i+1: self.status_var.set(f"正在处理表格 {num}..."))
                
                # 提取表格数据
                data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    data.append(row_data)
                
                # 写入Excel
                df = pd.DataFrame(data)
                sheet_name = f"表格{i+1}"
                df.to_excel(excel_writer, sheet_name=sheet_name, index=False, header=False)
                
                # 美化Excel表格
                self._beautify_excel_sheet(excel_writer, sheet_name, df)
            
            # 保存Excel文件
            excel_writer.close()
            
            # 完成提示
            self.root.after(0, lambda: messagebox.showinfo("成功", f"成功提取 {total_tables} 个表格到 {excel_path}"))
            
        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("错误", f"提取过程中出错: {str(e)}"))
        finally:
            # 重置UI
            self.root.after(0, self.reset_ui)
    
    def update_progress(self, value):
        self.progress["value"] = value
    
    def reset_ui(self):
        self.extract_button.config(state=tk.NORMAL)
        self.status_var.set("就绪")
        self.progress["value"] = 0
        
    def _beautify_excel_sheet(self, excel_writer, sheet_name, df):
        """美化Excel工作表"""
        # 获取工作表
        ws = excel_writer.book[sheet_name]
        
        # 定义样式
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        header_font = Font(bold=True, color='FFFFFF')
        header_fill = PatternFill(start_color='4F81BD', end_color='4F81BD', fill_type='solid')
        default_font = Font(name='微软雅黑', size=10)
        
        # 计算每列的最佳宽度
        column_widths = {}
        for col in df.columns:
            # 获取该列最长文本的长度
            max_length = 0
            for cell in df[col]:
                if isinstance(cell, str):
                    length = len(cell)
                else:
                    length = len(str(cell))
                if length > max_length:
                    max_length = length
            # 留一些余量
            column_widths[col] = min(max_length + 2, 50)  # 最大宽度限制为50
        
        # 设置第一行为表头样式
        for cell in ws[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = thin_border
        
        # 应用样式到所有单元格并设置列宽
        for row in ws.iter_rows(min_row=1, max_row=len(df)+1, min_col=1, max_col=len(df.columns)):
            for cell in row:
                # 跳过第一行（已经单独处理）
                if cell.row > 1:
                    cell.font = default_font
                    cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                    cell.border = thin_border
                
                # 设置列宽
                col_idx = cell.column - 1  # 转换为0-based索引
                if col_idx in column_widths:
                    ws.column_dimensions[get_column_letter(cell.column)].width = column_widths[col_idx]
        
        # 自动调整行高
        for row in ws.iter_rows(min_row=1, max_row=len(df)+1):
            ws.row_dimensions[row[0].row].height = None  # 重置行高以便自动调整

if __name__ == "__main__":
    root = tk.Tk()
    app = WordTableExtractor(root)
    root.mainloop()